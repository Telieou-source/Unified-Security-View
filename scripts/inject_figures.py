"""
inject_figures.py — Replace [INSERT FIGURE N: ...] placeholders in a .docx
with the corresponding screenshot images.

Placeholders live inside single-cell tables, not body paragraphs.
Some placeholders have their text split across multiple runs by Word's
spell-check markers (<w:proofErr>), so we concatenate all run text to match.

Usage:
  python3 scripts/inject_figures.py
"""

import re, shutil, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = os.path.dirname(os.path.abspath(__file__))
ROOT   = os.path.dirname(BASE)

SRC_DOC = os.path.join(ROOT, "attached_assets",
    "Unit_7_Assignment_\u2013_MSIT_5910-01_Capstone_1779454758016.docx")
OUT_DOC = os.path.join(ROOT, "Unit7_Assignment_With_Figures.docx")

SCREENSHOTS = {
    1: os.path.join(ROOT, "screenshots", "Unit7_Figure1_Typecheck.jpg"),
    2: os.path.join(ROOT, "screenshots", "Unit7_Figure2_BrowserConsole.jpg"),
    3: os.path.join(ROOT, "screenshots", "Unit7_Figure3_TC03AlertLog.jpg"),
    4: os.path.join(ROOT, "attached_assets", "screenshots",
                    "github_com_Telieou-source_Unified-Security-View_releases.png"),
    5: os.path.join(ROOT, "screenshots", "Unit7_Figure5_GitHubIssues.jpg"),
    6: os.path.join(ROOT, "screenshots", "Unit7_Figure6_BuildOutput.jpg"),
    7: os.path.join(ROOT, "screenshots", "Unit7_Figure7_PnpmAudit.jpg"),
    8: os.path.join(ROOT, "attached_assets", "screenshots",
                    "unified-security-view_replit_app.png"),
}

FIG_WIDTH = Inches(5.8)   # fits within 1-inch margins on US Letter

PLACEHOLDER_RE = re.compile(r'\[INSERT FIGURE (\d+):', re.IGNORECASE)


def para_full_text(para):
    """Concatenate all run texts, ignoring spell-check splits."""
    return "".join(r.text for r in para.runs)


def replace_para_with_image(para, fig_num, img_path):
    """
    Clear the table-cell paragraph and replace with a centred image,
    then append a small italic caption below.
    """
    original_text = para_full_text(para).strip() or para.text.strip()
    parent_cell_elem = para._element.getparent()  # <w:tc>

    # ── 1. Clear all existing runs and proofErr marks from the paragraph ──
    # Remove every child except <w:pPr>
    pPr = para._element.find(qn('w:pPr'))
    for child in list(para._element):
        if child.tag != qn('w:pPr'):
            para._element.remove(child)

    # Set paragraph alignment to centre
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'center')

    # ── 2. Add image run ──
    run = para.add_run()
    try:
        run.add_picture(img_path, width=FIG_WIDTH)
        print(f"  ✓ Figure {fig_num} inserted ({os.path.basename(img_path)})")
    except Exception as exc:
        run.text = f"[IMAGE MISSING: Figure {fig_num}]"
        print(f"  ✗ Figure {fig_num} FAILED: {exc}")
        return

    # ── 3. Append caption paragraph inside the same cell ──
    cap_p = OxmlElement('w:p')

    cap_pPr = OxmlElement('w:pPr')
    cap_jc  = OxmlElement('w:jc')
    cap_jc.set(qn('w:val'), 'center')
    cap_pPr.append(cap_jc)
    cap_p.append(cap_pPr)

    cap_r   = OxmlElement('w:r')
    cap_rPr = OxmlElement('w:rPr')

    # Italic
    cap_i = OxmlElement('w:i'); cap_rPr.append(cap_i)
    cap_ics = OxmlElement('w:iCs'); cap_rPr.append(cap_ics)

    # 9 pt font
    cap_sz  = OxmlElement('w:sz');  cap_sz.set(qn('w:val'), '18'); cap_rPr.append(cap_sz)
    cap_szc = OxmlElement('w:szCs'); cap_szc.set(qn('w:val'), '18'); cap_rPr.append(cap_szc)

    # Dark grey colour
    cap_col = OxmlElement('w:color'); cap_col.set(qn('w:val'), '595959'); cap_rPr.append(cap_col)

    cap_r.append(cap_rPr)

    cap_t = OxmlElement('w:t')
    cap_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    cap_t.text = original_text
    cap_r.append(cap_t)
    cap_p.append(cap_r)

    # Insert caption after the image paragraph inside the cell
    para._element.addnext(cap_p)


def main():
    print(f"Source : {SRC_DOC}")
    print(f"Output : {OUT_DOC}\n")

    # Verify all image files exist before we touch the docx
    missing = []
    for n, p in SCREENSHOTS.items():
        if not os.path.exists(p):
            missing.append((n, p))
    if missing:
        for n, p in missing:
            print(f"  ✗ Figure {n} image not found: {p}")
        print("\nAborting — fix missing images first.")
        return

    shutil.copy2(SRC_DOC, OUT_DOC)
    doc = Document(OUT_DOC)

    replacements = 0

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = para_full_text(para) or para.text
                    m = PLACEHOLDER_RE.search(text)
                    if m:
                        fig_num = int(m.group(1))
                        img_path = SCREENSHOTS[fig_num]
                        print(f"→ table[{ti}] Figure {fig_num} — replacing placeholder...")
                        replace_para_with_image(para, fig_num, img_path)
                        replacements += 1

    doc.save(OUT_DOC)
    print(f"\nDone — {replacements} figure(s) inserted → {OUT_DOC}")


if __name__ == "__main__":
    main()
