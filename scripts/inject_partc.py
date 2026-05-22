"""
inject_partc.py — Replace [INSERT FIGURE N:] and [INSERT CODE EXCERPT N:]
placeholders in Unit 7 Part C report with the corresponding screenshots.

Placeholders appear inside single-cell tables (same pattern as Unit 7 assignment).
The Appendix code excerpt inserts two images (CE5 clone/install + CE6 build/deploy).

Usage:
  python3 scripts/inject_partc.py
"""

import re, shutil, os
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE)

SRC_DOC = os.path.join(ROOT, "attached_assets",
    "Unit_7_Part_C_\u2013_Capstone_Project_Report_High-Side_Multi-Class_1779460759233.docx")
OUT_DOC = os.path.join(ROOT, "Unit7_PartC_With_Figures.docx")

SS = os.path.join(ROOT, "screenshots")
AS = os.path.join(ROOT, "attached_assets", "screenshots")

# ── Image mapping ─────────────────────────────────────────────────────────────
FIGURES = {
    1:  os.path.join(SS, "Figure1_ArchitectureDiagram.jpg"),
    2:  os.path.join(SS, "Figure8_UnifiedDashboard_allsix.jpg"),
    3:  os.path.join(SS, "Figure2_EventGeneratorLive.jpg"),
    4:  os.path.join(SS, "Figure3_CrossDomainGuard.jpg"),
    5:  os.path.join(SS, "Figure4_CorrelationAlert.jpg"),
    6:  os.path.join(SS, "Unit7_Figure1_Typecheck.jpg"),
    7:  os.path.join(SS, "Unit7_Figure3_TC03AlertLog.jpg"),
    8:  os.path.join(SS, "Figure6_HealthMonitor_60s.jpg"),
    9:  os.path.join(AS, "github_com_Telieou-source_Unified-Security-View_releases.png"),
    10: os.path.join(SS, "Figure10_ExecutiveSummary.jpg"),
}

# Each entry: list of image paths (Appendix has two stacked)
CODE_EXCERPTS = {
    1: [os.path.join(SS, "CodeExcerpt1_interface_definitions.jpg")],
    2: [os.path.join(SS, "CodeExcerpt2_sanitize_method.jpg")],
    3: [os.path.join(SS, "CodeExcerpt3_computeConfidence.jpg")],
    4: [os.path.join(SS, "CodeExcerpt4_runSimulationTick.jpg")],
    "appendix": [
        os.path.join(SS, "CodeExcerpt5_CloneInstall.jpg"),
        os.path.join(SS, "CodeExcerpt6_BuildDeploy.jpg"),
    ],
}

FIG_WIDTH  = Inches(5.8)
CODE_WIDTH = Inches(5.5)

FIGURE_RE   = re.compile(r'\[INSERT FIGURE (\d+):\s*(.*?)(?:\]|$)', re.IGNORECASE | re.DOTALL)
CODE_NUM_RE = re.compile(r'\[INSERT CODE EXCERPT (\d+):\s*(.*?)(?:\]|$)', re.IGNORECASE | re.DOTALL)
CODE_APP_RE = re.compile(r'\[INSERT CODE EXCERPT\s*\(Appendix\):\s*(.*?)(?:\]|$)', re.IGNORECASE | re.DOTALL)


# ── Caption builder ───────────────────────────────────────────────────────────

def _build_caption(prefix, description):
    """Return a clean readable caption string (no [INSERT ...] syntax)."""
    desc = description.strip().rstrip(']').strip()
    # Truncate very long descriptions to ~120 chars
    if len(desc) > 120:
        desc = desc[:117].rsplit(' ', 1)[0] + '…'
    return f"{prefix}. {desc}" if desc else prefix


# ── XML helpers ───────────────────────────────────────────────────────────────

def para_full_text(para):
    return "".join(r.text for r in para.runs)


def _set_centered(elem):
    pPr = elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        elem.insert(0, pPr)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'center')


def _clear_runs(elem):
    for child in list(elem):
        if child.tag != qn('w:pPr'):
            elem.remove(child)


def _make_caption_elem(caption_text):
    """Centred, italic, 9 pt, dark-grey <w:p> element."""
    p = OxmlElement('w:p')

    pPr = OxmlElement('w:pPr')
    jc  = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc); p.append(pPr)

    r   = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    for tag in ['w:i', 'w:iCs']:
        rPr.append(OxmlElement(tag))
    for tag, val in [('w:sz', '18'), ('w:szCs', '18')]:
        el = OxmlElement(tag); el.set(qn('w:val'), val); rPr.append(el)
    col = OxmlElement('w:color'); col.set(qn('w:val'), '595959'); rPr.append(col)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = caption_text
    r.append(t); p.append(r)
    return p


def _make_image_para_elem(doc, img_path, width):
    """
    Build a centred <w:p> with an inline picture by using a detached paragraph.
    """
    tmp = doc.add_paragraph()
    _set_centered(tmp._element)
    tmp.add_run().add_picture(img_path, width=width)
    tmp._element.getparent().remove(tmp._element)
    return tmp._element


# ── Core injector ─────────────────────────────────────────────────────────────

def inject(doc, para, img_paths, width, label, caption_text):
    """
    Replace placeholder paragraph content with image(s) + caption.
    Extra images (Appendix) are inserted as sibling paragraphs after the first.
    """
    _clear_runs(para._element)
    _set_centered(para._element)

    # First image: inject directly into the cleared placeholder paragraph
    run = para.add_run()
    try:
        run.add_picture(img_paths[0], width=width)
        print(f"  ✓ {label} image 1 — {os.path.basename(img_paths[0])}")
    except Exception as exc:
        run.text = f"[IMAGE ERROR: {label}]"
        print(f"  ✗ {label} image 1 FAILED: {exc}")
        return

    insert_after = para._element

    # Extra images (e.g. Appendix CE5 + CE6)
    for i, extra_path in enumerate(img_paths[1:], start=2):
        try:
            extra_elem = _make_image_para_elem(doc, extra_path, width)
            insert_after.addnext(extra_elem)
            insert_after = extra_elem
            print(f"  ✓ {label} image {i} — {os.path.basename(extra_path)}")
        except Exception as exc:
            print(f"  ✗ {label} image {i} FAILED: {exc}")

    # Caption (clean, no [INSERT ...] syntax)
    cap = _make_caption_elem(caption_text)
    insert_after.addnext(cap)


# ── Dispatcher ────────────────────────────────────────────────────────────────

def process_para(doc, para):
    text = para_full_text(para) or para.text
    if not text or '[INSERT' not in text.upper():
        return False

    # Appendix code excerpt — check before numbered CE pattern
    m = CODE_APP_RE.search(text)
    if m:
        caption = _build_caption("Appendix B — Setup Commands", m.group(1))
        print(f"→ Code Excerpt (Appendix)")
        inject(doc, para, CODE_EXCERPTS["appendix"], CODE_WIDTH,
               "Code Excerpt (Appendix)", caption)
        return True

    # Numbered code excerpt
    m = CODE_NUM_RE.search(text)
    if m:
        n = int(m.group(1))
        if n in CODE_EXCERPTS:
            caption = _build_caption(f"Code Excerpt {n}", m.group(2))
            print(f"→ Code Excerpt {n}")
            inject(doc, para, CODE_EXCERPTS[n], CODE_WIDTH,
                   f"Code Excerpt {n}", caption)
            return True

    # Figure
    m = FIGURE_RE.search(text)
    if m:
        n = int(m.group(1))
        if n in FIGURES:
            caption = _build_caption(f"Figure {n}", m.group(2))
            print(f"→ Figure {n}")
            inject(doc, para, [FIGURES[n]], FIG_WIDTH, f"Figure {n}", caption)
            return True

    return False


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print(f"Source : {SRC_DOC}")
    print(f"Output : {OUT_DOC}\n")

    missing = []
    for n, p in FIGURES.items():
        if not os.path.exists(p):
            missing.append(f"Figure {n}: {p}")
    for k, paths in CODE_EXCERPTS.items():
        for p in paths:
            if not os.path.exists(p):
                missing.append(f"CE {k}: {p}")
    if missing:
        for m in missing: print(f"  ✗ {m}")
        print("\nAborting — fix missing images first.")
        return

    shutil.copy2(SRC_DOC, OUT_DOC)
    doc = Document(OUT_DOC)

    total = 0

    print("── Body paragraphs ──")
    for para in doc.paragraphs:
        if process_para(doc, para):
            total += 1

    print("\n── Table cells ──")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if process_para(doc, para):
                        total += 1

    doc.save(OUT_DOC)
    size = os.path.getsize(OUT_DOC)
    print(f"\nDone — {total} placeholder(s) replaced")
    print(f"Output : {OUT_DOC}  ({size:,} bytes  /  {size // 1024:,} KB)")


if __name__ == "__main__":
    main()
